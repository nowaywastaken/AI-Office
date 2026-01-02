from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, List, Dict, Any
from io import BytesIO
import os


class WordEngine:
    """
    A comprehensive Word document generator with full formatting control.
    Supports: fonts, sizes, colors, spacing, margins, tables, images, and more.
    """
    
    def __init__(self):
        self.doc = Document()
        self._setup_default_styles()
    
    def _setup_default_styles(self):
        """Set up default document styles."""
        # Set default font for the document
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        # For Chinese font support
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ==================== PAGE SETUP ====================
    
    def set_page_margins(
        self,
        top: float = 2.54,
        bottom: float = 2.54,
        left: float = 3.18,
        right: float = 3.18,
        unit: str = 'cm'
    ):
        """
        Set page margins.
        
        Args:
            top, bottom, left, right: Margin values
            unit: 'cm' or 'inches'
        """
        sections = self.doc.sections
        for section in sections:
            if unit == 'cm':
                section.top_margin = Cm(top)
                section.bottom_margin = Cm(bottom)
                section.left_margin = Cm(left)
                section.right_margin = Cm(right)
            else:
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
        return self
    
    def set_page_size(self, width: float, height: float, unit: str = 'cm'):
        """Set page size (A4 default: 21.0 x 29.7 cm)."""
        for section in self.doc.sections:
            if unit == 'cm':
                section.page_width = Cm(width)
                section.page_height = Cm(height)
            else:
                section.page_width = Inches(width)
                section.page_height = Inches(height)
        return self
    
    # ==================== PARAGRAPH OPERATIONS ====================
    
    def add_paragraph(
        self,
        text: str,
        font_name: Optional[str] = None,
        font_size: Optional[float] = None,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        color: Optional[str] = None,
        alignment: str = 'left',
        line_spacing: Optional[float] = None,
        line_spacing_rule: str = 'multiple',
        space_before: Optional[float] = None,
        space_after: Optional[float] = None,
        first_line_indent: Optional[float] = None
    ):
        """
        Add a paragraph with full formatting control.
        
        Args:
            text: The paragraph text
            font_name: Font family (e.g., 'Arial', 'Times New Roman', '微软雅黑')
            font_size: Font size in points
            bold, italic, underline: Text decoration
            color: Hex color code (e.g., '#FF0000')
            alignment: 'left', 'center', 'right', 'justify'
            line_spacing: Line spacing value
            line_spacing_rule: 'single', 'double', '1.5', 'multiple', 'exact', 'at_least'
            space_before: Spacing before paragraph in points
            space_after: Spacing after paragraph in points
            first_line_indent: First line indent in cm
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        
        # Font settings
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = Pt(font_size)
        
        run.font.bold = bold
        run.font.italic = italic
        run.font.underline = underline
        
        if color:
            run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))
        
        # Paragraph alignment
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Line spacing
        pf = paragraph.paragraph_format
        if line_spacing:
            if line_spacing_rule == 'single':
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif line_spacing_rule == 'double':
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            elif line_spacing_rule == '1.5':
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif line_spacing_rule == 'exact':
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(line_spacing)
            elif line_spacing_rule == 'at_least':
                pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                pf.line_spacing = Pt(line_spacing)
            else:  # multiple
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = line_spacing
        
        # Paragraph spacing
        if space_before is not None:
            pf.space_before = Pt(space_before)
        if space_after is not None:
            pf.space_after = Pt(space_after)
        
        # First line indent
        if first_line_indent is not None:
            pf.first_line_indent = Cm(first_line_indent)
        
        return paragraph
    
    def add_heading(
        self,
        text: str,
        level: int = 1,
        font_name: Optional[str] = None,
        font_size: Optional[float] = None,
        color: Optional[str] = None
    ):
        """Add a heading with optional custom formatting."""
        heading = self.doc.add_heading(text, level=level)
        
        if font_name or font_size or color:
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if font_size:
                    run.font.size = Pt(font_size)
                if color:
                    run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))
        
        return heading
    
    # ==================== TABLE OPERATIONS ====================
    
    def add_table(
        self,
        data: List[List[str]],
        header: bool = True,
        style: str = 'Table Grid',
        col_widths: Optional[List[float]] = None
    ):
        """
        Add a table with data.
        
        Args:
            data: 2D list of strings representing table data
            header: If True, first row is treated as header
            style: Table style name
            col_widths: List of column widths in cm
        """
        if not data:
            return None
        
        rows = len(data)
        cols = len(data[0]) if data else 0
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style
        
        # Set column widths if provided
        if col_widths:
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Cm(width)
        
        # Populate table
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = str(cell_text)
                # Bold header row
                if header and i == 0:
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        
        return table
    
    # ==================== IMAGE OPERATIONS ====================
    
    def add_image(
        self,
        image_path: str,
        width: Optional[float] = None,
        height: Optional[float] = None,
        unit: str = 'cm'
    ):
        """Add an image to the document."""
        if width and unit == 'cm':
            width = Cm(width)
        elif width and unit == 'inches':
            width = Inches(width)
        
        if height and unit == 'cm':
            height = Cm(height)
        elif height and unit == 'inches':
            height = Inches(height)
        
        return self.doc.add_picture(image_path, width=width, height=height)
    
    # ==================== SAVE OPERATIONS ====================
    
    def save(self, filepath: str):
        """Save document to file."""
        self.doc.save(filepath)
        return filepath
    
    def save_to_bytes(self) -> bytes:
        """Save document to bytes (for streaming)."""
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def create_blank(self):
        """Return the underlying document object."""
        return self.doc
